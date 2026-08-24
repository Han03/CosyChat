"""文档文本提取工具，支持 txt/md/pdf/docx 等格式。"""

import io
import os

from utils.logger import logger


def extract_document_text(file_bytes: bytes, filename: str, max_chars: int = 300) -> str:
    """从文档文件中提取文本内容。

    Args:
        file_bytes: 文件二进制内容
        filename: 文件名（用于判断扩展名）
        max_chars: 最大提取字符数，防止超长文档拖慢模型推理

    Returns:
        提取的文本内容；失败时返回空字符串。
    """
    ext = os.path.splitext(filename)[1].lower()
    try:
        if ext in (".txt", ".md"):
            text = _extract_text_plain(file_bytes)
        elif ext == ".pdf":
            text = _extract_text_pdf(file_bytes)
        elif ext == ".docx":
            text = _extract_text_docx(file_bytes)
        elif ext == ".doc":
            text = _extract_text_doc(file_bytes)
        else:
            logger.warning(f"[DocParser] 不支持的文档格式: {ext}")
            return ""
    except Exception as e:
        logger.error(f"[DocParser] 提取文档 '{filename}' 文本失败: {e}")
        return ""

    if not text:
        return ""

    text = text.strip()
    if len(text) > max_chars:
        text = text[:max_chars] + "\n...(文档内容过长，已截断)"
        logger.info(f"[DocParser] 文档 '{filename}' 内容已截断至 {max_chars} 字符")
    logger.info(f"[DocParser] 从 '{filename}' 提取文本，长度: {len(text)}")
    return text


def _extract_text_plain(file_bytes: bytes) -> str:
    """提取 txt/md 文件文本，自动尝试常见编码。"""
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def _extract_text_pdf(file_bytes: bytes) -> str:
    """使用 pdfplumber 提取 PDF 文本。"""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_text_docx(file_bytes: bytes) -> str:
    """使用 python-docx 提取 .docx 文本。"""
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text:
            text_parts.append(para.text)
    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def _extract_text_doc(file_bytes: bytes) -> str:
    """.doc（旧版二进制格式）需要 antiword 或 win32com，此处降级提示。"""
    logger.warning("[DocParser] .doc 旧格式不支持，请转换为 .docx 后上传")
    return ""
